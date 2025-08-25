from langchain_openai import AzureChatOpenAI
from langgraph.graph import StateGraph, START, END
from typing import Annotated
from langchain_core.messages import AnyMessage, SystemMessage, HumanMessage
from langgraph.graph.message import add_messages
from operator import add
from docx import Document
from src.components.replace_field_text import (
    replace_text,
)
from langgraph.graph import START, StateGraph
from langgraph.prebuilt import ToolNode
from typing import TypedDict
import io
from io import BytesIO

from src.components.azure_auth import (
    get_token_provider_default,
    get_token_provider_streamlit_secrets,
)

# -------- AZURE AUTHENTICATION --------
token_provider = get_token_provider_streamlit_secrets()
# If you want to use Streamlit secrets authentication, use:
# token_provider = get_token_provider_streamlit_secrets()

llm = AzureChatOpenAI(
    azure_endpoint="https://oai02-aiserv.openai.azure.com/",
    api_version="2024-10-21",
    azure_ad_token_provider=token_provider,
    # azure_deployment="gpt-4.1-nano",
    # azure_deployment="gpt-4.1",
    azure_deployment="gpt-4o-2024-08-06",
    temperature=0.1,
)

tools = [replace_text]
llm_with_tools = llm.bind_tools(tools, tool_choice="replace_text")
# System message
sys_msg = SystemMessage(
    content="Du er en hjælpsom assistent der finder passager i dokumenter, og vurderer hvad de skal erstattes med, baseret på brugerens input. Brug værktøjerne til at hjælpe med dette."
)


class OverallState(TypedDict):
    messages: Annotated[list[AnyMessage], add_messages]
    document: Annotated[list[bytes], add]


# Node
def assistant(state: OverallState):
    for m in state["messages"]:
        m.pretty_print()
    return {"messages": [llm_with_tools.invoke([sys_msg] + state["messages"])]}


# Graph
builder = StateGraph(OverallState)

# Define nodes: these do the work
builder.add_node("assistant", assistant)
builder.add_node("tools", ToolNode(tools))

# Define edges: these determine how the control flow moves
builder.add_edge(START, "assistant")
builder.add_edge("assistant", "tools")
builder.add_edge("tools", END)

react_graph = builder.compile()


def start_graph_llm(user_prompt: str, document_bytes: bytes):
    print(f"\n STARTING GRAPH LLM PROCESSING...")
    doc = Document(io.BytesIO(document_bytes))
    document_text = "\n".join([para.text for para in doc.paragraphs])

    initial_message = (
        "\n---BRUGERPROMPT:---\n"
        + user_prompt
        + "\n---DOKUMENT-TEKST---\n"
        + document_text
    )
    messages = [HumanMessage(content=initial_message)]
    output = react_graph.invoke(
        {"messages": messages, "document": [document_bytes]}, {"recursion_limit": 5}
    )
    # Output progression of all documents
    for m in output["messages"]:
        m.pretty_print()
    for idx, doc_bytes in enumerate(output["document"]):
        doc = Document(BytesIO(doc_bytes))
        doc_text = "\n".join([para.text for para in doc.paragraphs])
        print(
            f"\n--- DOCUMENT {idx+1}/{len(output['document'])} ---\n{doc_text}\n--- END DOCUMENT {idx+1} ---\n"
        )

    return output


def start_graph_llm_fake(user_prompt: str, document_bytes: bytes):
    print(f"\n STARTING FAKE GRAPH LLM PROCESSING...")

    from langchain_core.messages import HumanMessage, AIMessage
    from docx import Document
    from io import BytesIO

    # Simulate extracting document text
    doc = Document(BytesIO(document_bytes))
    document_text = "\n".join([para.text for para in doc.paragraphs])

    # Create fake messages
    initial_message = (
        "\n---BRUGERPROMPT:---\n"
        + user_prompt
        + "\n---DOKUMENT-TEKST---\n"
        + document_text
    )
    messages = [
        HumanMessage(content=initial_message),
        AIMessage(
            content="Dette er et simuleret svar fra agenten. Ingen ændringer er foretaget."
        ),
    ]
    output = {"messages": messages, "document": [document_bytes]}

    return output
