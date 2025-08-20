from src.components.llm import graph
from langchain_core.messages import HumanMessage, SystemMessage
from docx import Document
import json
import re
from loguru import logger
import traceback


def process_document_with_llm(document: Document):
    # Extract all text from the docx document
    text = "\n".join([para.text for para in document.paragraphs])
    # print("Old document text:\n\n", text)
    # Prepare messages for LLM
    request = {"messages": [sys_msg] + [HumanMessage(content=text)]}
    # Get response from LLM
    messages = graph.invoke(request)
    for m in messages["messages"]:
        m.pretty_print()
    print("\n\n--- LLM Response ---\n")
    print("LLM response:\n\n", messages["messages"][-1].content)
    return messages["messages"][-1].content


sys_msg = SystemMessage(
    content="""Du modtager nu en tekst, hvor du får følgende opgave: Du vil se et mønster hvor der står "if betingelse" (case insensitive), <en arbitrær mængde ord - lad os kalde dem MIDTERORD>, og så på et tidspunkt vil der stå ”<et antal tegn>” Else ”<et antal tegn>”. Altså 2 citationstegn med noget indeni, og så 2 sitationstegn mere, med noget andet indeni, lad os kalde dem TEKST1 og TEKST2. Dette skal du transformere til følgende: { IF "J" { MERGEFIELD <MIDTERORD>}" " TEKST1" " TEKST2" Du skal kun ændre noget i teksten ved dette specifikke mønster. Returnér i json-format det der erstattes + det det erstattes med for hver case. Eksempel på json-format: 
    {
        "replaced": "if betingelse CustomerName ... \"Hej\" Else \"Farvel\"",
        "with": "{ IF \"J\" { MERGEFIELD CustomerName } \"Hej\" \"Farvel\" }"
    }"""
)


def process_document_with_llm_fake(document: Document):
    """
    Returns a mock object with a .content attribute containing a JSON string,
    matching the expected format for apply_llm_replacements.
    """

    class FakeLLMResponse:
        def __init__(self, content):
            self.content = content

    replacements = [
        {
            "replaced": "If betingelse Borger enlig ved ældrecheck berettigelse”din” Else ”jeres”",
            "with": '{ IF "J" = "{ MERGEFIELD Borger enlig ved ældrecheck berettigelse }" " din" " jeres" }',
        },
        {
            "replaced": "If betingelse Borger enlig ved ældrecheck berettigelse  ”din” Else ”jeres”",
            "with": '{ IF "J" = "{ MERGEFIELD Borger enlig ved ældrecheck berettigelse }" " din" " jeres" }',
        },
    ]
    return FakeLLMResponse(json.dumps(replacements))


def replace_titles_with_keys(mapping: dict, document: Document):
    """
    Replace occurrences of each 'Titel' in the document with its corresponding 'Nøgle'.
    Starts with the longest 'Titel' to avoid partial replacements.
    """
    # Sort keys by length descending to avoid partial replacements
    sorted_titles = sorted(mapping.keys(), key=len, reverse=True)
    for para in document.paragraphs:
        text = para.text
        for titel in sorted_titles:
            nøgle = mapping[titel]
            text = text.replace(titel, nøgle)
        para.text = text
    return document


def apply_llm_replacements(llm_response_message, document: Document):
    """
    Apply replacements from LLM response to the document.
    Handles text that may span across multiple paragraphs and runs.
    """
    try:
        # Parse the JSON response from LLM
        response_content = llm_response_message.content

        # Extract JSON array or object from the response (in case there's extra text)
        json_match = re.search(r"(\[.*\]|\{.*\})", response_content, re.DOTALL)
        if not json_match:
            print("No JSON found in LLM response")
            return document

        replacements_data = json.loads(json_match.group())

        # Handle both single replacement and list of replacements
        if isinstance(replacements_data, dict) and "replaced" in replacements_data:
            replacements_data = [replacements_data]
        elif isinstance(replacements_data, list):
            pass
        else:
            print("Unexpected JSON format")
            return document

        # Apply each replacement
        for replacement in replacements_data:
            replaced_text = replacement.get("replaced", "")
            with_text = replacement.get("with", "")

            if replaced_text and with_text:
                document = _replace_text_in_document(document, replaced_text, with_text)

    except json.JSONDecodeError as e:
        error_type = type(e).__name__
        tb = traceback.format_exc()
        logger.debug(
            f"Fejl ved behandling af Word-dokument:\n\n"
            f"Type: {error_type}\n"
            f"Detaljer: {str(e)}\n\n"
            f"Traceback:\n{tb}"
        )
    except Exception as e:
        error_type = type(e).__name__
        tb = traceback.format_exc()
        logger.debug(
            f"Fejl ved behandling af Word-dokument:\n\n"
            f"Type: {error_type}\n"
            f"Detaljer: {str(e)}\n\n"
            f"Traceback:\n{tb}"
        )

    return document


def _replace_text_in_document(document: Document, search_text: str, replace_text: str):
    """
    Replace text in document, handling text that spans across paragraphs and runs.
    """
    # Normalize whitespace for searching
    search_normalized = re.sub(r"\s+", " ", search_text.strip())

    # Get all text content with position tracking
    all_text = ""
    text_map = []  # Track (paragraph_index, run_index, start_pos, end_pos)

    for para_idx, paragraph in enumerate(document.paragraphs):
        para_start = len(all_text)

        for run_idx, run in enumerate(paragraph.runs):
            run_start = len(all_text)
            run_text = run.text
            all_text += run_text
            run_end = len(all_text)

            if run_text:  # Only add non-empty runs
                text_map.append((para_idx, run_idx, run_start, run_end))

        # Add paragraph break (except for last paragraph)
        if para_idx < len(document.paragraphs) - 1:
            all_text += " "

    # Normalize the full text for searching
    all_text_normalized = re.sub(r"\s+", " ", all_text)

    # Find the search text in normalized content
    search_pos = all_text_normalized.find(search_normalized)
    if search_pos == -1:
        print(f"Text not found: {search_text[:50]}...")
        return document

    search_end = search_pos + len(search_normalized)

    # Find which runs are affected
    affected_runs = []
    for para_idx, run_idx, start_pos, end_pos in text_map:
        # Check if this run overlaps with the search range
        if start_pos < search_end and end_pos > search_pos:
            affected_runs.append((para_idx, run_idx, start_pos, end_pos))

    if not affected_runs:
        return document

    # Perform the replacement
    # Clear affected runs first
    for para_idx, run_idx, _, _ in affected_runs:
        document.paragraphs[para_idx].runs[run_idx].text = ""

    # Add the replacement text to the first affected run
    if affected_runs:
        first_para_idx, first_run_idx, _, _ = affected_runs[0]
        document.paragraphs[first_para_idx].runs[first_run_idx].text = replace_text

    return document
    return document
