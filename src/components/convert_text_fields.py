from docx import Document
import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import logging

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


class FieldConverter:
    """Class to convert text representations of Word fields into actual Word fields."""

    def __init__(self):
        self.field_regex = re.compile(r"{([^{}]*)({[^{}]*})[^{}]*}|{([^{}]*)}")

    def process_document(self, doc):
        """Process a Word document and convert text fields to actual fields."""
        logger.info("Processing document in-place")

        # Process all paragraphs in the document
        for paragraph in doc.paragraphs:
            self._process_paragraph(paragraph)

        # Process all tables in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_paragraph(paragraph)

        logger.info("Document processing complete")
        return doc

    def _process_paragraph(self, paragraph):
        """Process a paragraph to find and convert text fields."""
        # Store original text for debugging and comparison
        original_text = paragraph.text

        # Check if the paragraph might contain field text
        if "{" in original_text and "}" in original_text:
            logger.debug(f"Processing paragraph: {original_text[:50]}...")
            self._convert_paragraph_fields(paragraph)

    def _convert_paragraph_fields(self, paragraph):
        """Convert text fields in a paragraph to actual fields."""
        # We need to work with runs since fields can span multiple runs
        runs_text = "".join([run.text for run in paragraph.runs])

        # Find potential fields in the combined text
        field_matches = self._find_field_matches(runs_text)

        if field_matches:
            # Clear the paragraph
            for i in range(len(paragraph.runs) - 1, -1, -1):
                paragraph._p.remove(paragraph.runs[i]._r)

            # Add the text back with actual fields
            self._add_text_with_fields(paragraph, runs_text, field_matches)

    def _find_field_matches(self, text):
        """Find all potential field matches in text."""
        # This is a simplified version - a real implementation would need
        # more sophisticated parsing to handle nested structures correctly
        matches = []

        # Process the text to find field patterns
        start_pos = 0
        while True:
            match = self._find_next_field(text, start_pos)
            if not match:
                break

            matches.append(match)
            start_pos = match["end"]

        return matches

    def _find_next_field(self, text, start_pos=0):
        """Find the next field in the text starting from start_pos."""
        # This needs to handle nested fields and is quite complex
        # A simplified approach for demonstration
        open_braces = 0
        field_start = -1

        for i in range(start_pos, len(text)):
            if text[i] == "{":
                if open_braces == 0:
                    field_start = i
                open_braces += 1
            elif text[i] == "}":
                open_braces -= 1
                if open_braces == 0 and field_start != -1:
                    field_text = text[field_start : i + 1]
                    field_type, params = self._parse_field(field_text)

                    if field_type:
                        return {
                            "start": field_start,
                            "end": i + 1,
                            "text": field_text,
                            "type": field_type,
                            "params": params,
                        }
                    break

        return None

    def _parse_field(self, field_text):
        """Parse a field text to extract its type and parameters."""
        # Remove outer braces and trim
        inner_text = field_text[1:-1].strip()

        # Check for IF field
        if inner_text.startswith("IF "):
            parts = self._split_if_field(inner_text[3:])
            return "IF", parts

        # Check for MERGEFIELD
        if inner_text.startswith("MERGEFIELD "):
            field_name = inner_text[11:].strip()
            return "MERGEFIELD", [field_name]

        # Add other field types as needed

        return None, []

    def _split_if_field(self, if_text):
        """Split an IF field text into its components."""
        # This is complex due to nested fields
        # A simplified version for demonstration
        parts = []
        current_part = ""
        in_quotes = False
        open_braces = 0

        for char in if_text:
            if char == '"' and open_braces == 0:
                in_quotes = not in_quotes
                current_part += char
            elif char == "{":
                open_braces += 1
                current_part += char
            elif char == "}":
                open_braces -= 1
                current_part += char
            elif char in (" ") and not in_quotes and open_braces == 0:
                if current_part:
                    parts.append(current_part.strip())
                    current_part = ""
            else:
                current_part += char

        if current_part:
            parts.append(current_part.strip())

        return parts

    def _add_text_with_fields(self, paragraph, original_text, field_matches):
        """Add text back to paragraph with actual fields replacing text fields."""
        last_end = 0

        for match in field_matches:
            # Add text before this field
            if match["start"] > last_end:
                paragraph.add_run(original_text[last_end : match["start"]])

            # Add the actual field
            self._add_field(paragraph, match)

            last_end = match["end"]

        # Add any remaining text
        if last_end < len(original_text):
            paragraph.add_run(original_text[last_end:])

    def _add_field(self, paragraph, field_match):
        """Add an actual Word field to the paragraph."""
        if field_match["type"] == "MERGEFIELD":
            self._add_merge_field(paragraph, field_match["params"][0])
        elif field_match["type"] == "IF":
            self._add_if_field(paragraph, field_match["params"])
        # Add other field types as needed

    def _add_merge_field(self, paragraph, field_name):
        """Add a MERGEFIELD to the paragraph."""
        run = paragraph.add_run()

        # Create the field begin element
        begin = self._create_element("w:fldChar", {"w:fldCharType": "begin"})
        run._r.append(begin)

        # Create the field code
        instr_text = self._create_element("w:instrText", {})
        instr_text.text = f" MERGEFIELD {field_name} "
        run._r.append(instr_text)

        # Create the field separator
        separator = self._create_element("w:fldChar", {"w:fldCharType": "separate"})
        run._r.append(separator)

        # Create the field end element
        end = self._create_element("w:fldChar", {"w:fldCharType": "end"})
        run._r.append(end)

    def _add_if_field(self, paragraph, if_params):
        """Add an IF field to the paragraph."""
        # This is complex due to nested fields and multiple parameters
        # A simplified implementation for demonstration
        run = paragraph.add_run()

        # Create the field begin element
        begin = self._create_element("w:fldChar", {"w:fldCharType": "begin"})
        run._r.append(begin)

        # Create the field code for IF
        instr_text = self._create_element("w:instrText", {})
        instr_text.text = f' IF {" ".join(if_params)} '
        run._r.append(instr_text)

        # Create the field separator
        separator = self._create_element("w:fldChar", {"w:fldCharType": "separate"})
        run._r.append(separator)

        # Add result text here if needed

        # Create the field end element
        end = self._create_element("w:fldChar", {"w:fldCharType": "end"})
        run._r.append(end)

    def _create_element(self, name, attrs=None):
        """Create an XML element with namespace."""
        element = OxmlElement(name)
        if attrs:
            for key, value in attrs.items():
                element.set(qn(key), value)
        return element


def convert_document_fields(document):
    """Convert text fields in a document to actual Word fields and return the modified document."""
    converter = FieldConverter()
    return converter.process_document(document)


if __name__ == "__main__":
    # Example usage
    import sys

    if len(sys.argv) != 3:
        print("Usage: python convert_text_fields.py <input_docx> <output_docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    doc = convert_document_fields(input_file)
    doc.save(output_file)
    print(f"Document processed and saved to {output_file}")
