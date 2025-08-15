from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re
from .find_fields import replace_titles_with_mergefields_respecting_colors


def load_docx(file_like):
    """Load a Word document from a file-like object."""
    document = Document(file_like)

    return document


def save_docx_to_bytes(doc):
    """Save a Document object to BytesIO."""
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


def extract_colors_from_paragraph(paragraph):
    """
    Extracts all unique text and background colors from a paragraph.

    Args:
        paragraph (Paragraph): A `Paragraph` object from `python-docx`.

    Returns:
        tuple: A tuple containing two lists:
            - List of unique text colors.
            - List of unique background colors.
    """
    text_colors = set()
    background_colors = set()

    for run in paragraph.runs:
        # Extract text color
        if run.font.color and run.font.color.rgb:
            text_colors.add(str(run.font.color.rgb))

        # Extract background color from highlight elements
        highlights = run._element.xpath(".//w:highlight")
        for highlight in highlights:
            bg_color = highlight.get(qn("w:val"))
            if bg_color:
                background_colors.add(bg_color)

        # Also check for shading elements (alternative way highlights can be stored)
        shadings = run._element.xpath(".//w:shd")
        for shading in shadings:
            # Check for fill color
            fill_color = shading.get(qn("w:fill"))
            if fill_color and fill_color != "auto":
                background_colors.add(fill_color)

            # Check for color attribute
            shd_color = shading.get(qn("w:color"))
            if shd_color and shd_color != "auto":
                background_colors.add(shd_color)

    return list(text_colors), list(background_colors)


def process_paragraph_with_color_aware_replacements(paragraph, mappings: dict):
    """
    Process a paragraph to replace titles with MERGEFIELD keys, respecting color consistency.

    Args:
        paragraph (Paragraph): A `Paragraph` object from `python-docx`.
        mappings (dict): Dictionary mapping titles to their corresponding keys.

    Returns:
        str: The modified paragraph text.
    """
    # Replace titles respecting color consistency
    modified_text = replace_titles_with_mergefields_respecting_colors(
        paragraph, mappings
    )

    # Update the paragraph text
    paragraph.text = modified_text

    # Insert actual merge fields
    return paragraph.text


def create_mergefield_element(field_name):
    """
    Create a Word MERGEFIELD element with the given field name.

    Args:
        field_name (str): The name of the merge field

    Returns:
        OxmlElement: A complete mergefield element ready to be inserted
    """
    # Create the field begin element
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    # Create the instruction text element
    instrText = OxmlElement("w:instrText")
    instrText.text = f"MERGEFIELD {field_name}"

    # Create the field end element
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    return fldChar_begin, instrText, fldChar_end


def insert_actual_mergefields(paragraph):
    """
    Convert text-based MERGEFIELD syntax to actual Word merge fields in a paragraph.

    This function finds all occurrences of { MERGEFIELD field_name } in the paragraph
    and replaces them with actual Word merge field elements that can be toggled with Alt+F9.

    Args:
        paragraph (Paragraph): A Word paragraph object from python-docx
    """
    # Pattern to match { MERGEFIELD field_name } with required spaces
    pattern = r"\{\s+MERGEFIELD\s+([^}]+)\s+\}"

    # Get the paragraph element
    p_element = paragraph._element

    # We need to work with runs to properly insert the merge fields
    # First, let's rebuild the paragraph by processing each run

    # Collect all text and run information
    runs_info = []
    for run in paragraph.runs:
        runs_info.append({"text": run.text, "element": run._element, "font": run.font})

    # Clear all runs
    for run in paragraph.runs:
        p_element.remove(run._element)

    # Process the concatenated text to find merge fields
    full_text = "".join([r["text"] for r in runs_info])

    # Find all merge field matches
    matches = list(re.finditer(pattern, full_text, re.IGNORECASE))

    if not matches:
        # No merge fields found, restore original runs
        for run_info in runs_info:
            p_element.append(run_info["element"])
        return

    # Process text and insert merge fields
    last_end = 0
    current_run_idx = 0
    current_run_pos = 0

    for match in matches:
        start_pos = match.start()
        end_pos = match.end()
        field_name = match.group(1).strip()

        # Add text before the merge field
        if start_pos > last_end:
            text_before = full_text[last_end:start_pos]
            _insert_text_with_runs(
                paragraph, text_before, runs_info, current_run_idx, current_run_pos
            )
            # Update position tracking
            text_len = len(text_before)
            current_run_idx, current_run_pos = _advance_position(
                runs_info, current_run_idx, current_run_pos, text_len
            )

        # Insert the merge field
        _insert_mergefield(paragraph, field_name)

        # Update position tracking past the merge field text
        field_text_len = end_pos - start_pos
        current_run_idx, current_run_pos = _advance_position(
            runs_info, current_run_idx, current_run_pos, field_text_len
        )

        last_end = end_pos

    # Add any remaining text after the last merge field
    if last_end < len(full_text):
        remaining_text = full_text[last_end:]
        _insert_text_with_runs(
            paragraph, remaining_text, runs_info, current_run_idx, current_run_pos
        )


def _insert_text_with_runs(paragraph, text, runs_info, start_run_idx, start_run_pos):
    """Helper function to insert text while preserving run formatting."""
    if not text:
        return

    remaining_text = text
    run_idx = start_run_idx
    run_pos = start_run_pos

    while remaining_text and run_idx < len(runs_info):
        run_info = runs_info[run_idx]
        available_in_run = len(run_info["text"]) - run_pos

        if available_in_run <= 0:
            run_idx += 1
            run_pos = 0
            continue

        # Take what we can from this run
        take_length = min(len(remaining_text), available_in_run)
        text_to_insert = remaining_text[:take_length]

        # Create a new run with the original formatting
        new_run = paragraph.add_run(text_to_insert)

        # Copy formatting from original run if available
        if run_idx < len(runs_info):
            original_run_element = runs_info[run_idx]["element"]
            # Copy run properties if they exist
            rPr = original_run_element.find(qn("w:rPr"))
            if rPr is not None:
                new_rPr = OxmlElement("w:rPr")
                new_rPr[:] = rPr[:]
                new_run._element.insert(0, new_rPr)

        remaining_text = remaining_text[take_length:]
        run_pos += take_length

        if run_pos >= len(run_info["text"]):
            run_idx += 1
            run_pos = 0


def _insert_mergefield(paragraph, field_name):
    """Helper function to insert an actual merge field."""
    # Create three runs for the merge field: begin, instruction, end

    # Begin run
    begin_run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    begin_run._element.append(fldChar_begin)

    # Instruction run
    instr_run = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    # Add a space before and after the MERGEFIELD content to match the curly bracket spacing
    instrText.text = f" MERGEFIELD {field_name} "
    instr_run._element.append(instrText)

    # End run
    end_run = paragraph.add_run()
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    end_run._element.append(fldChar_end)


def _advance_position(runs_info, run_idx, run_pos, text_length):
    """Helper function to advance position tracking through runs."""
    remaining = text_length

    while remaining > 0 and run_idx < len(runs_info):
        available_in_run = len(runs_info[run_idx]["text"]) - run_pos

        if available_in_run <= remaining:
            remaining -= available_in_run
            run_idx += 1
            run_pos = 0
        else:
            run_pos += remaining
            remaining = 0

    return run_idx, run_pos
