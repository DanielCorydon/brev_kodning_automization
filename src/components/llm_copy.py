from azure.identity import DefaultAzureCredential, get_bearer_token_provider
from langchain_openai import AzureChatOpenAI
from langgraph.graph import StateGraph, START, END
from pydantic import BaseModel, Field
from typing import Annotated
from langchain_core.messages import AnyMessage, SystemMessage, HumanMessage
from langgraph.graph.message import add_messages
from operator import add
from docx import Document
from src.components.replace_field_text import (
    replace_text_from_json,
)
from src.components.test_tools import replace_text_from_json_test

import json
from typing import TypedDict
from langgraph.graph import MessagesState
import io
from io import BytesIO


# -------- AZURE AUTHENTICATION --------
credential = DefaultAzureCredential(
    exclude_environment_credential=True,
    exclude_developer_cli_credential=True,
    exclude_workload_identity_credential=True,
    exclude_managed_identity_credential=True,
    exclude_visual_studio_code_credential=True,
    exclude_shared_token_cache_credential=True,
    exclude_interactive_browser_credential=True,
)
token_provider = get_bearer_token_provider(
    credential, "https://cognitiveservices.azure.com/.default"
)


llm = AzureChatOpenAI(
    azure_endpoint="https://oai02-aiserv.openai.azure.com/",
    api_version="2024-10-21",
    azure_ad_token_provider=token_provider,
    # azure_deployment="gpt-4.1-nano",
    # azure_deployment="gpt-4.1",
    azure_deployment="gpt-4o-2024-08-06",
    temperature=0.1,
)

from langchain_openai import ChatOpenAI


tools = [replace_text_from_json_test]
llm_with_tools = llm.bind_tools(tools)
# System message
sys_msg = SystemMessage(
    content="You are a helpful assistant tasked with performing arithmetic on a set of inputs, or NLP tasks on documents."
)


class OverallState(TypedDict):
    messages: Annotated[list[AnyMessage], add_messages]
    document: Annotated[list[bytes], add]


# Node
def assistant(state: OverallState):
    for m in state["messages"]:
        m.pretty_print()
    return {"messages": [llm_with_tools.invoke([sys_msg] + state["messages"])]}


from langgraph.graph import START, StateGraph
from langgraph.prebuilt import tools_condition, ToolNode

# Graph
builder = StateGraph(OverallState)

# Define nodes: these do the work
builder.add_node("assistant", assistant)
builder.add_node("tools", ToolNode(tools))

# Define edges: these determine how the control flow moves
builder.add_edge(START, "assistant")
builder.add_conditional_edges(
    "assistant",
    # If the latest message (result) from assistant is a tool call -> tools_condition routes to tools
    # If the latest message (result) from assistant is a not a tool call -> tools_condition routes to END
    tools_condition,
)
builder.add_edge("tools", "assistant")
react_graph = builder.compile()

# messages = [HumanMessage(content="Add 3 and 4.")]
# messages = react_graph.invoke({"messages": messages})
# for m in messages["messages"]:
#     m.pretty_print()


def start_graph_llm(user_prompt: str, document_bytes: bytes):
    doc = Document(io.BytesIO(document_bytes))
    document_text = "\n".join([para.text for para in doc.paragraphs])

    initial_message = (
        "\n---BRUGERPROMPT:---\n"
        + user_prompt
        + "\n---DOKUMENT-TEKST---\n"
        + document_text
    )
    messages = [HumanMessage(content=initial_message)]
    output = react_graph.invoke(
        {"messages": messages, "document": [document_bytes]}, {"recursion_limit": 5}
    )
    # Output progression of all documents
    for idx, doc_bytes in enumerate(output["document"]):
        doc = Document(BytesIO(doc_bytes))
        doc_text = "\n".join([para.text for para in doc.paragraphs])
        print(
            f"\n--- DOCUMENT {idx+1}/{len(output['document'])} ---\n{doc_text}\n--- END DOCUMENT {idx+1} ---\n"
        )
    for m in output["messages"]:
        m.pretty_print()

    return messages
