from docx import Document
from io import BytesIO
import json
import re
import logging
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from langchain_core.tools import tool
from typing import IO
from typing import Annotated, NotRequired
from langchain_core.tools import tool
from langgraph.prebuilt import InjectedState, create_react_agent
from langgraph.prebuilt.chat_agent_executor import AgentState
from operator import add
from typing import Annotated
from langgraph.types import Command
from langchain_core.messages import ToolMessage
from langchain_core.tools import tool, InjectedToolCallId
from langchain_core.messages import AnyMessage, SystemMessage, HumanMessage
from langgraph.graph.message import add_messages


class CustomState(AgentState):
    # The user_name field in short-term state
    messages: Annotated[list[AnyMessage], add_messages]
    document: Annotated[list[bytes], add]


logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


@tool
def replace_text_from_json_test(
    state: Annotated[CustomState, InjectedState],
    replace: str,
    replacement: str,
    tool_call_id: Annotated[str, InjectedToolCallId],
) -> Command:
    """
    Replaces all occurrences of 'replace' with 'replacement' in a Word document.
    Word document is being passed via state, so should be ignored by LLM agent.

    Args:
        state: InjectedState containing a list of document bytes. The last document is used.
        replace: The text to replace.
        replacement: The text to replace with.

    Returns:
        The modified document as bytes (docx format).
    """
    doc = Document(BytesIO(state.get("document")[-1]))
    for paragraph in doc.paragraphs:
        if replace:
            paragraph.text = paragraph.text.replace(replace, replacement)
    print("Replaced text in document ", doc.paragraphs[0].text)
    output_stream = BytesIO()
    doc.save(output_stream)

    return Command(
        update={
            "messages": [
                ToolMessage(f"Replaced text in document", tool_call_id=tool_call_id)
            ],
            "document": [output_stream.getvalue()],
        }
    )
