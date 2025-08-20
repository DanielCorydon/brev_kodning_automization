import re
from docx import Document
from docx.enum.text import WD_COLOR  # Added for highlight color

PARA_SEP = (
    "\u2029"  # paragraph separator to ensure regexes can cross paragraphs unambiguously
)


def _iter_paragraph_runs(doc):
    """
    Yield (p_idx, r_idx, run) for each run in body paragraphs.
    (Extend to headers/footers/tables if needed.)
    """
    for p_idx, p in enumerate(doc.paragraphs):
        for r_idx, run in enumerate(p.runs):
            yield p_idx, r_idx, run


def _build_text_index(doc):
    """
    Build:
      - flat_text: str with PARA_SEP between paragraphs
      - index: list of dicts, one per character in flat_text:
          { 'p_idx': int or None, 'r_idx': int or None, 'run': Run or None, 'offset_in_run': int or None, 'char': str }
        (For PARA_SEP chars, run fields are None.)
    Also returns:
      - para_starts: list of start indices in flat_text for each paragraph (for debugging/optional use)
    """
    index = []
    parts = []
    para_starts = []
    for p_idx, p in enumerate(doc.paragraphs):
        para_starts.append(len(parts))
        for r_idx, run in enumerate(p.runs):
            t = run.text or ""
            for i, ch in enumerate(t):
                index.append(
                    {
                        "p_idx": p_idx,
                        "r_idx": r_idx,
                        "run": run,
                        "offset_in_run": i,
                        "char": ch,
                    }
                )
                parts.append(ch)
        # Insert a paragraph separator *between* paragraphs (not after the last one)
        if p_idx < len(doc.paragraphs) - 1:
            index.append(
                {
                    "p_idx": p_idx,
                    "r_idx": None,
                    "run": None,
                    "offset_in_run": None,
                    "char": PARA_SEP,
                }
            )
            parts.append(PARA_SEP)

    return "".join(parts), index, para_starts


def _collect_affected_run_spans(index, start, end):
    """
    Given a match [start, end) in flat_text, return an ordered list of affected runs with
    exact span slices per run:
      [ { 'run': Run, 'p_idx': int, 'r_idx': int, 'start_off': int, 'end_off': int } ... ]
    PARA_SEP chars are ignored; only real run chars are included.
    """
    spans = []
    i = start
    current = None

    while i < end:
        cell = index[i]
        i += 1
        if cell["run"] is None:  # paragraph separator or other non-run char
            # finalize any ongoing span
            if current is not None:
                spans.append(current)
                current = None
            continue

        r = cell["run"]
        # start a new span if run changed
        if (current is None) or (current["run"] is not r):
            if current is not None:
                spans.append(current)
            current = {
                "run": r,
                "p_idx": cell["p_idx"],
                "r_idx": cell["r_idx"],
                "start_off": cell["offset_in_run"],
                "end_off": cell["offset_in_run"] + 1,
            }
        else:
            # extend current span
            current["end_off"] = cell["offset_in_run"] + 1

    if current is not None:
        spans.append(current)

    # order by (p_idx, r_idx) to ensure deterministic forward order in the document
    spans.sort(key=lambda s: (s["p_idx"], s["r_idx"]))
    return spans


def _remove_run(run):
    # Delete a run entirely from its paragraph
    r_el = run._element
    p_el = r_el.getparent()
    p_el.remove(r_el)


def _replace_in_runs(spans, replacement_text):
    """
    Apply one replacement across a list of run spans (ordered).
    - First span's run: replace covered segment with replacement_text (keep style).
    - Middle spans: remove runs if fully covered; otherwise trim text portion.
    - Last span: if same as first (single-run match), it's already handled.
    Returns the first run where replacement was written.
    """
    if not spans:
        return None

    # First affected run
    first = spans[0]
    first_run = first["run"]
    f_text = first_run.text or ""
    # Replacement goes into the exact slice in the first run
    first_left = f_text[: first["start_off"]]
    first_right = f_text[first["end_off"] :]
    first_run.text = first_left + replacement_text + first_right

    # Highlight the replaced text in green
    first_run.font.highlight_color = WD_COLOR.GREEN

    # Subsequent spans
    for i, s in enumerate(spans[1:], start=1):
        run = s["run"]
        t = run.text or ""
        # Decide if we should remove or trim
        covered_len = s["end_off"] - s["start_off"]
        if s["start_off"] == 0 and s["end_off"] == len(t):
            # Entire run covered by the match -> remove it
            _remove_run(run)
        else:
            # Partial coverage -> cut away the matched slice
            new_text = t[: s["start_off"]] + t[s["end_off"] :]
            run.text = new_text

    return first_run


def regex_replace_docx(
    doc: Document,
    pattern: str,
    repl,
    flags=re.DOTALL,
    allow_newlines_in_replacement=False,
):
    """
    Perform a regex replacement across a docx Document, spanning paragraphs and runs.
    Replacement text is inserted into the *first affected run* while preserving its style.
    Any following runs covered by the match are removed (or trimmed if partially covered).

    Args:
        doc: python-docx Document
        pattern: regex pattern (use DOTALL to span paragraphs; DOTALL is default)
        repl: str or callable(match) -> str
        flags: regex flags (default DOTALL)
        allow_newlines_in_replacement: if False, any newlines in replacement are converted to spaces
                                       (recommended; creating new paragraphs programmatically is trickier)

    Returns:
        count of replacements performed
    """
    flat, index, _ = _build_text_index(doc)
    regex = re.compile(pattern, flags)

    # Collect matches first, then process from the end
    matches = list(regex.finditer(flat))
    if not matches:
        return 0

    # Process in reverse so earlier spans’ character indexing stays valid
    # We will resolve spans using the frozen 'index' character map we built up-front.
    replaced = 0
    for m in reversed(matches):
        start, end = m.span()
        spans = _collect_affected_run_spans(index, start, end)
        if not spans:
            continue

        # Compute replacement text
        if callable(repl):
            rtext = repl(m)
        else:
            # Support backreferences like \1
            rtext = m.expand(repl)

        if not allow_newlines_in_replacement and ("\n" in rtext or "\r" in rtext):
            rtext = re.sub(r"[\r\n]+", " ", rtext)

        _replace_in_runs(spans, rtext)
        replaced += 1

    return replaced


# ------------------------
# Example usage:
# ------------------------
# doc = Document("input.docx")
# # Example: replace {{SOMETHING … ACROSS PARAGRAPHS}} with "[REPLACED]"
# n = regex_replace_docx(doc, r"\{\{.*?\}\}", "[REPLACED]")
# print(f"Replaced {n} matches")
# doc.save("output.docx")
