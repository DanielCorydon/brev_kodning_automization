from docx import Document
from io import BytesIO
import json
import re
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from langchain_core.tools import tool
from typing import IO
from typing import Annotated, NotRequired
from langchain_core.tools import tool
from langgraph.prebuilt import InjectedState, create_react_agent
from langgraph.prebuilt.chat_agent_executor import AgentState
from operator import add
from typing import Annotated
from langgraph.types import Command
from langchain_core.messages import ToolMessage
from langchain_core.tools import tool, InjectedToolCallId
from langchain_core.messages import AnyMessage, SystemMessage, HumanMessage
from langgraph.graph.message import add_messages
from loguru import logger


class CustomState(AgentState):
    # The user_name field in short-term state
    messages: Annotated[list[AnyMessage], add_messages]
    document: Annotated[list[bytes], add]


class CustomState(AgentState):
    # The user_name field in short-term state
    document: Annotated[list[bytes], add]


@dataclass
class ReplacementMatch:
    """Data class to store information about a text match and its replacement."""

    original_text: str
    replacement_text: str
    start_pos: int
    end_pos: int


@tool
def replace_text(
    state: Annotated[CustomState, InjectedState],
    tool_call_id: Annotated[str, InjectedToolCallId],
    replacement_pairs: List[Dict],
) -> Document:
    """
    Replace text in document based on replacement pairs.

    Args:
        replacement_pairs: List of dictionaries, each with:
            - "originalText": The text to search for and replace.
            - "replacementText": The text to use as the replacement.

    Returns:
        Modified document object
    """

    # Load the document from bytes
    doc = Document(BytesIO(state.get("document")[-1]))

    # Process all paragraphs in the document
    for paragraph in doc.paragraphs:
        _process_paragraph(paragraph, replacement_pairs)

    # Process all tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _process_paragraph(paragraph, replacement_pairs)
    output_stream = BytesIO()
    doc.save(output_stream)
    logger.info("Text replacement process complete")
    return Command(
        update={
            "messages": [
                ToolMessage(
                    f"Alle tekstbidder er nu blevet erstattet i dokumentet.",
                    tool_call_id=tool_call_id,
                )
            ],
            "document": [output_stream.getvalue()],
        }
    )


def _process_paragraph(paragraph, json_data: List[Dict]):
    """Process a single paragraph for text replacements."""
    if not paragraph.runs:
        return

    # Get the complete text from all runs in the paragraph
    full_paragraph_text = _get_paragraph_text(paragraph)

    if not full_paragraph_text.strip():
        return

    logger.debug(f"Processing paragraph: '{full_paragraph_text[:100]}...'")

    # Find all matches for all patterns in this paragraph
    all_matches = _find_all_matches(full_paragraph_text, json_data)

    if all_matches:
        logger.debug(f"Found {len(all_matches)} matches in paragraph")

        # Remove overlapping matches (keep the first occurrence)
        all_matches = _remove_overlapping_matches(all_matches)
        logger.debug(f"After removing overlaps: {len(all_matches)} matches")

        # Sort matches by position (reverse order to replace from end to beginning)
        all_matches.sort(key=lambda x: x.start_pos, reverse=True)

        # Apply replacements one by one
        for i, match in enumerate(all_matches):
            logger.debug(
                f"Applying match {i+1}/{len(all_matches)}: '{match.original_text}' -> '{match.replacement_text}'"
            )
            _apply_replacement(paragraph, match)


def _get_paragraph_text(paragraph) -> str:
    """Get the complete text from all runs in a paragraph."""
    return "".join(run.text for run in paragraph.runs)


def _find_all_matches(
    paragraph_text: str, json_data: List[Dict]
) -> List[ReplacementMatch]:
    """Find all matches for all patterns in the given text."""
    matches = []

    # Iterate through each match object in the JSON data
    for replacement_data in json_data:
        original_text = replacement_data.get("originalText", "")
        replacement_text = replacement_data.get("replacementText", "")

        if not original_text:
            continue

        # Find all occurrences of the originalText in the paragraph
        text_matches = _find_text_occurrences(paragraph_text, original_text)

        for start_pos, end_pos in text_matches:
            match = ReplacementMatch(
                original_text=original_text,
                replacement_text=replacement_text,
                start_pos=start_pos,
                end_pos=end_pos,
            )
            matches.append(match)

    return matches


def _find_text_occurrences(text: str, search_text: str) -> List[Tuple[int, int]]:
    """Find all occurrences of search_text in text, handling case sensitivity and whitespace."""
    matches = []

    # First try exact matching
    exact_matches = _find_exact_matches(text, search_text)
    if exact_matches:
        return exact_matches

    # If no exact matches, try normalized matching
    # Clean up the search text - normalize quotes and whitespace
    search_text_clean = _normalize_text_for_matching(search_text)
    text_clean = _normalize_text_for_matching(text)

    if not search_text_clean:
        return matches

    logger.debug(f"Looking for: '{search_text_clean}' in: '{text_clean}'")

    # Use case-insensitive search to handle variations
    start = 0
    while True:
        pos = text_clean.lower().find(search_text_clean.lower(), start)
        if pos == -1:
            break

        logger.debug(f"Found match at normalized position {pos}")

        # Map back to original text positions
        original_start = _map_normalized_to_original_position(text, text_clean, pos)
        original_end = _map_normalized_to_original_position(
            text, text_clean, pos + len(search_text_clean)
        )

        # Validate that this is a reasonable match by checking the actual text
        if original_end <= len(text):
            actual_text = text[original_start:original_end]
            # Only accept if the match looks reasonable (similar length and content)
            if _is_reasonable_match(actual_text, search_text):
                logger.debug(
                    f"Mapped to original positions: {original_start}-{original_end}"
                )
                logger.debug(f"Original text segment: '{actual_text}'")
                matches.append((original_start, original_end))

        start = pos + len(search_text_clean)  # Move past this match

    return matches


def _find_exact_matches(text: str, search_text: str) -> List[Tuple[int, int]]:
    """Find exact matches of search_text in text."""
    matches = []
    start = 0

    while True:
        pos = text.find(search_text, start)
        if pos == -1:
            break
        matches.append((pos, pos + len(search_text)))
        start = pos + 1

    return matches


def _is_reasonable_match(found_text: str, search_text: str) -> bool:
    """Check if the found text is a reasonable match for the search text."""
    # Check length similarity (within 20% difference)
    len_ratio = len(found_text) / len(search_text) if len(search_text) > 0 else 0
    if len_ratio < 0.8 or len_ratio > 1.2:
        return False

    # Check if key words are present
    search_words = search_text.lower().split()
    found_words = found_text.lower().split()

    # At least 70% of words should be present
    common_words = sum(1 for word in search_words if word in found_words)
    word_ratio = common_words / len(search_words) if len(search_words) > 0 else 0

    return word_ratio >= 0.7


def _normalize_text_for_matching(text: str) -> str:
    """Normalize text for better matching by handling quotes and whitespace."""
    # Replace different types of quotes with standard quotes
    text = re.sub(r'["""]', '"', text)
    text = re.sub(r"[''']", "'", text)

    # Normalize whitespace but preserve structure
    text = re.sub(r"\s+", " ", text.strip())

    return text


def _map_normalized_to_original_position(
    original_text: str, normalized_text: str, norm_pos: int
) -> int:
    """Map a position in normalized text back to original text."""
    if norm_pos <= 0:
        return 0
    if norm_pos >= len(normalized_text):
        return len(original_text)

    # Build character-by-character mapping
    orig_to_norm = []  # Maps original position to normalized position
    norm_pos_current = 0

    i = 0
    while i < len(original_text):
        orig_char = original_text[i]
        orig_to_norm.append(norm_pos_current)

        # Handle quote normalization
        if orig_char in '"' "'":
            norm_pos_current += 1  # Normalized to standard quote
        elif orig_char.isspace():
            # Skip multiple consecutive whitespace in original
            while i < len(original_text) and original_text[i].isspace():
                orig_to_norm.append(norm_pos_current)
                i += 1
            norm_pos_current += 1  # Single space in normalized
            continue  # i already incremented
        else:
            norm_pos_current += 1

        i += 1

    orig_to_norm.append(norm_pos_current)

    # Find the original position that maps to the desired normalized position
    for orig_pos, mapped_norm_pos in enumerate(orig_to_norm):
        if mapped_norm_pos >= norm_pos:
            return orig_pos

    return len(original_text)


def _find_actual_match_position(
    current_text: str, search_text: str, expected_start: int
) -> Optional[Tuple[int, int]]:
    """Find the actual position of the search text in the current text."""
    # First try the expected position
    expected_end = expected_start + len(search_text)
    if (
        expected_start < len(current_text)
        and expected_end <= len(current_text)
        and current_text[expected_start:expected_end] == search_text
    ):
        return (expected_start, expected_end)

    # If that fails, search for the text in the vicinity
    search_start = max(0, expected_start - 50)
    search_end = min(len(current_text), expected_start + len(search_text) + 50)
    search_area = current_text[search_start:search_end]

    # Normalize both texts for comparison
    normalized_search = _normalize_text_for_matching(search_text)
    normalized_area = _normalize_text_for_matching(search_area)

    # Look for the normalized text
    pos = normalized_area.lower().find(normalized_search.lower())
    if pos != -1:
        # Map back to original positions
        actual_start = search_start + _map_normalized_to_original_position(
            search_area, normalized_area, pos
        )
        actual_end = search_start + _map_normalized_to_original_position(
            search_area, normalized_area, pos + len(normalized_search)
        )
        return (actual_start, actual_end)

    # Last resort: search the entire text
    normalized_current = _normalize_text_for_matching(current_text)
    pos = normalized_current.lower().find(normalized_search.lower())
    if pos != -1:
        actual_start = _map_normalized_to_original_position(
            current_text, normalized_current, pos
        )
        actual_end = _map_normalized_to_original_position(
            current_text, normalized_current, pos + len(normalized_search)
        )
        return (actual_start, actual_end)

    return None


def _fix_replacement_text(replacement_text: str) -> str:
    """Fix common issues in replacement text format."""
    text = replacement_text.strip()

    # Handle the case where the replacement is already properly formatted
    if text.startswith("{ IF ") and text.endswith(" }"):
        return text

    # Ensure proper Word field format: { IF "J" = { MERGEFIELD field-name }" " text1" " text2" }
    if text.startswith('{ IF "J" = { MERGEFIELD ') and not text.endswith(" }"):
        # Add missing closing brace
        text += " }"

    # Fix common quote and spacing issues
    # Ensure proper spacing around the MERGEFIELD closing brace
    text = re.sub(r'}\s*"\s*"\s*', r'}" "', text)

    # Fix spacing around the text values
    text = re.sub(r'"\s*"\s*([^"]+)"\s*"\s*', r'" "\1" "', text)

    return text


def _remove_overlapping_matches(
    matches: List[ReplacementMatch],
) -> List[ReplacementMatch]:
    """Remove overlapping matches, keeping the first (longest/most specific) one."""
    if not matches:
        return matches

    # Sort by start position first
    sorted_matches = sorted(matches, key=lambda x: x.start_pos)
    non_overlapping = []

    for match in sorted_matches:
        # Check if this match overlaps with any accepted match
        overlaps = False
        for accepted in non_overlapping:
            if (
                match.start_pos < accepted.end_pos
                and match.end_pos > accepted.start_pos
            ):
                overlaps = True
                break

        if not overlaps:
            non_overlapping.append(match)

    return non_overlapping


def _apply_replacement(paragraph, match: ReplacementMatch):
    """Apply a single text replacement while preserving styling."""
    # Get current paragraph text (it may have changed from previous replacements)
    current_text = _get_paragraph_text(paragraph)

    # For replacements done in reverse order, we need to search for the actual text
    # instead of relying on fixed positions
    actual_match_pos = _find_actual_match_position(
        current_text, match.original_text, match.start_pos
    )

    if actual_match_pos is None:
        logger.warning(
            f"Could not find text '{match.original_text}' in current paragraph text"
        )
        return

    actual_start, actual_end = actual_match_pos

    # Verify the match is still valid
    if actual_end > len(current_text):
        logger.warning(
            f"Match position {actual_start}-{actual_end} exceeds current text length {len(current_text)}"
        )
        return

    # Get the styling from the first character of the match
    source_style = _get_style_at_position(paragraph, actual_start)

    # Fix the replacement text format
    fixed_replacement = _fix_replacement_text(match.replacement_text)

    # Create new text with replacement
    new_text = (
        current_text[:actual_start] + fixed_replacement + current_text[actual_end:]
    )

    logger.debug(
        f"Replacing '{current_text[actual_start:actual_end]}' with '{fixed_replacement}'"
    )

    # Rebuild paragraph with new text
    _replace_paragraph_text(
        paragraph,
        new_text,
        source_style,
        actual_start,
        len(fixed_replacement),
    )


def _get_style_at_position(paragraph, position: int) -> Dict:
    """Get the style of the run at the specified position."""
    current_pos = 0
    for run in paragraph.runs:
        run_end = current_pos + len(run.text)
        if current_pos <= position < run_end:
            return _extract_run_style(run)
        current_pos = run_end

    # If position is at the end, use the last run's style
    if paragraph.runs:
        return _extract_run_style(paragraph.runs[-1])

    return {}


def _replace_paragraph_text(
    paragraph,
    new_text: str,
    replacement_style: Dict,
    replacement_start: int,
    replacement_length: int,
):
    """Replace the entire paragraph text while preserving as much styling as possible."""
    # Store original run information
    original_runs = []
    current_pos = 0

    for run in paragraph.runs:
        run_info = {
            "text": run.text,
            "style": _extract_run_style(run),
            "start": current_pos,
            "end": current_pos + len(run.text),
        }
        original_runs.append(run_info)
        current_pos += len(run.text)

    # Clear all runs
    for i in range(len(paragraph.runs) - 1, -1, -1):
        paragraph._element.remove(paragraph.runs[i]._element)

    # Add new text with appropriate styling
    if not new_text:
        return

    replacement_end = replacement_start + replacement_length

    # Text before replacement
    if replacement_start > 0:
        before_text = new_text[:replacement_start]
        before_style = _find_style_for_original_position(original_runs, 0)
        if before_text:
            run = paragraph.add_run(before_text)
            _apply_run_style(run, before_style)

    # Replacement text
    if replacement_length > 0:
        replacement_text = new_text[replacement_start:replacement_end]
        if replacement_text:
            run = paragraph.add_run(replacement_text)
            _apply_run_style(run, replacement_style)

    # Text after replacement
    if replacement_end < len(new_text):
        after_text = new_text[replacement_end:]
        # The after text should use the style from the original position after the match
        after_style = _find_style_for_original_position(
            original_runs, replacement_start
        )
        if after_text:
            run = paragraph.add_run(after_text)
            _apply_run_style(run, after_style)


def _find_style_for_original_position(original_runs: List[Dict], position: int) -> Dict:
    """Find the style that was at a given position in the original text."""
    for run_info in original_runs:
        if run_info["start"] <= position < run_info["end"]:
            return run_info["style"]

    # If position is beyond the original text, use the last run's style
    if original_runs:
        return original_runs[-1]["style"]

    return {}


def _extract_run_style(run) -> Dict:
    """Extract styling information from a run."""
    style = {
        "font_name": run.font.name,
        "font_size": run.font.size,
        "bold": run.font.bold,
        "italic": run.font.italic,
        "underline": run.font.underline,
        "color": None,
        "highlight_color": run.font.highlight_color,
    }

    # Extract color if present
    if run.font.color.rgb:
        style["color"] = run.font.color.rgb

    return style


def _apply_run_style(run, style: Dict):
    """Apply styling information to a run."""
    if not style:
        return

    if style.get("font_name"):
        run.font.name = style["font_name"]
    if style.get("font_size"):
        run.font.size = style["font_size"]
    if style.get("bold") is not None:
        run.font.bold = style["bold"]
    if style.get("italic") is not None:
        run.font.italic = style["italic"]
    if style.get("underline") is not None:
        run.font.underline = style["underline"]
    if style.get("color"):
        run.font.color.rgb = style["color"]
    if style.get("highlight_color"):
        run.font.highlight_color = style["highlight_color"]


def process_document_from_json_file(
    doc_path: str, json_path: str, output_path: str
) -> Document:
    """
    Process a document using a JSON file containing replacement patterns.

    Args:
        doc_path: Path to the input Word document
        json_path: Path to the JSON file with replacement patterns
        output_path: Path where the modified document should be saved

    Returns:
        Modified document object
    """
    # Load the document
    doc = Document(doc_path)

    # Load the JSON data
    with open(json_path, "r", encoding="utf-8") as f:
        json_data = json.load(f)

    # Process the document
    modified_doc = replace_text(doc, json_data)

    # Save the modified document
    modified_doc.save(output_path)
    logger.info(f"Document saved to {output_path}")

    return modified_doc


def process_document_from_json_string(
    doc_path: str, json_string: str, output_path: str
) -> Document:
    """
    Process a document using a JSON string containing replacement patterns.

    Args:
        doc_path: Path to the input Word document
        json_string: JSON string with replacement patterns
        output_path: Path where the modified document should be saved

    Returns:
        Modified document object
    """
    # Load the document
    doc = Document(doc_path)

    # Parse the JSON data
    json_data = json.loads(json_string)

    # Process the document
    modified_doc = replace_text(doc, json_data)

    # Save the modified document
    modified_doc.save(output_path)
    logger.info(f"Document saved to {output_path}")

    return modified_doc


def replace_text_in_document(doc: Document, json_data: List[Dict]) -> Document:
    """
    Convenience function to replace text in a document based on JSON data.

    Args:
        doc: Word document object
        json_data: List of match dicts

    Returns:
        Modified document object
    """
    return replace_text(doc, json_data)


if __name__ == "__main__":
    # Example usage
    import sys

    if len(sys.argv) < 4:
        print(
            "Usage: python replace_field_text.py <input_docx> <json_file> <output_docx>"
        )
        print(
            "   or: python replace_field_text.py <input_docx> <json_string> <output_docx> --json-string"
        )
        sys.exit(1)

    input_file = sys.argv[1]
    json_input = sys.argv[2]
    output_file = sys.argv[3]

    if len(sys.argv) > 4 and sys.argv[4] == "--json-string":
        process_document_from_json_string(input_file, json_input, output_file)
    else:
        process_document_from_json_file(input_file, json_input, output_file)

    print(f"Document processed and saved to {output_file}")
