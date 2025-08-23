"""
Module for extracting and formatting regex matches from Word documents.
"""

import re
from docx import Document
import logging
from typing import Dict, List, Any, Pattern, Optional, Set
from io import BytesIO
import json  # Add this import to fix the NameError

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


class DocumentRegexFinder:
    """
    Finds and extracts text matching regex patterns from Word documents.
    """

    def __init__(self):
        """Initialize the DocumentRegexFinder."""
        pass

    def get_document_text(self, doc_input: Any) -> str:
        """
        Returns all text from a Word document, including paragraphs and tables, as a single string.

        Args:
            doc_input: Path to the Word document or a BytesIO object.

        Returns:
            str: All document text with paragraphs separated by newlines.
        """
        try:
            if isinstance(doc_input, str):
                doc = Document(doc_input)
            elif isinstance(doc_input, BytesIO):
                doc = Document(doc_input)
            else:
                raise ValueError(
                    "Invalid input type. Expected file path or BytesIO object."
                )

            # Extract text from paragraphs, preserving paragraph breaks
            paragraphs_text = [para.text for para in doc.paragraphs]

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs_text.extend([para.text for para in cell.paragraphs])

            # Join all paragraphs with newlines to maintain structure
            full_text = "\n".join(paragraphs_text)

            return full_text

        except Exception as e:
            logger.error(f"Error extracting text from document: {e}")
            raise

    def find_regex_matches_in_document(
        self, doc_path: str, patterns: List[Pattern]
    ) -> List[Dict[str, Any]]:
        """
        Finds all unique matches for each regex pattern in the document text.

        Args:
            doc_path: Path to the Word document.
            patterns: List of compiled regex patterns.

        Returns:
            list: List of match dicts, each with regex, fullText, groups.
        """
        doc_text = self.get_document_text(doc_path)
        results = []
        seen_full_texts: Set[str] = set()
        for pattern in patterns:
            pattern_str = pattern.pattern
            matches = pattern.finditer(doc_text)
            for match in matches:
                full_text = match.group(0)
                if full_text in seen_full_texts:
                    continue
                seen_full_texts.add(full_text)
                groups = [
                    match.group(i)
                    for i in range(1, len(match.groups()) + 1)
                    if match.group(i) is not None
                ]
                results.append(
                    {"regex": pattern_str, "fullText": full_text, "groups": groups}
                )
        return results


def extract_and_format_regex_matches(
    doc_path: str, regex_list: List[str]
) -> List[Dict[str, Any]]:
    """
    Extracts regex matches from a Word document.

    Args:
        doc_path: Path to the Word document.
        regex_list: List of regex pattern strings.

    Returns:
        list: List of match dicts.
    """
    compiled_patterns = []
    for regex_str in regex_list:
        try:
            pattern = re.compile(regex_str, re.DOTALL | re.UNICODE)
            compiled_patterns.append(pattern)
        except re.error as e:
            logger.error(f"Invalid regex pattern '{regex_str}': {e}")
    finder = DocumentRegexFinder()
    results = finder.find_regex_matches_in_document(doc_path, compiled_patterns)
    return results


if __name__ == "__main__":
    # Example usage
    import sys

    if len(sys.argv) < 3:
        print(
            "Usage: python find_change_sentences.py <docx_file> <regex1> [regex2] ..."
        )
        sys.exit(1)
    doc_path = sys.argv[1]
    regexes = sys.argv[2:]
    results = extract_and_format_regex_matches(doc_path, regexes)
    print("\n--- Processed Results ---\n")
    print(json.dumps(results, indent=2, ensure_ascii=False))
    doc_path = sys.argv[1]
    regexes = sys.argv[2:]
    results = extract_and_format_regex_matches(doc_path, regexes)
    print("\n--- Processed Results ---\n")
    print(json.dumps(results, indent=2, ensure_ascii=False))
