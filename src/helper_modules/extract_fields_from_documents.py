"""
Module for extracting Word fields from documents.

This module provides functionality to extract all Word fields from a document,
including nested fields, and display them in a readable format.

OVERVIEW:
=========
Word documents contain fields that are dynamic content placeholders. These fields
can be simple (like a MERGEFIELD) or complex with nested structures (like an IF
field containing MERGEFIELD fields). This module extracts and displays these
field structures in their complete form.

FIELD STRUCTURE:
===============
Word fields have three main components:
1. Field Begin: Marks the start of a field (fldCharType="begin")
2. Field Code: The instruction text (e.g., "MERGEFIELD Name")
3. Field Separator: Divides code from result (fldCharType="separate")
4. Field Result: The displayed value when field is evaluated
5. Field End: Marks the end of a field (fldCharType="end")

NESTED FIELDS:
=============
Fields can be nested within other fields. For example:
{ IF "condition" = { MERGEFIELD Value } "true text" "false text" }

The challenge is properly reconstructing the complete nested structure while
avoiding duplication of individual field components.

EXTRACTION LOGIC:
================
The extraction works by:
1. Tracking field nesting levels as we encounter begin/end markers
2. Building field codes by collecting instruction text and regular text
3. Properly handling nested field boundaries with curly braces
4. Only returning top-level fields (nesting level 1) to avoid duplicates
5. Filtering out result text when building field codes

OUTPUT FORMAT:
=============
Each extracted field contains:
- type: The field type (IF, MERGEFIELD, etc.)
- code: The complete field instruction
- result: The field's evaluated result text
- full_text: The complete field in { } format
- nested_fields: List of nested fields found within this field
"""

import sys
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
import sys
import re
from docx import Document
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import xml.etree.ElementTree as ET


class WordFieldExtractor:
    """
    A class to extract Word fields from documents, including nested fields.

    This class can handle various types of Word fields such as MERGEFIELD, IF, etc.
    and properly parse nested field structures. It uses the python-docx library
    to access the underlying XML structure of Word documents.

    Key Features:
    - Extracts complete nested field structures
    - Handles complex IF statements with multiple MERGEFIELD components
    - Preserves original field formatting and structure
    - Identifies field types and nested relationships
    - Provides both individual field codes and complete nested representations

    Processing Logic:
    - Parses XML runs sequentially to maintain field order
    - Tracks nesting levels using field begin/end markers
    - Separates field instruction text from result text
    - Reconstructs complete field codes with proper brace matching
    - Filters out field results to show only the instruction syntax

    Example Output:
    For a field like: { IF "condition" = { MERGEFIELD Value } "true" "false" }
    - Extracts the complete nested structure as one field
    - Identifies IF as the main field type
    - Lists { MERGEFIELD Value } as a nested field
    - Preserves exact formatting and structure
    """

    def __init__(self, document_path: str):
        """
        Initialize the extractor with a document path.

        Args:
            document_path (str): Path to the Word document
        """
        self.document_path = Path(document_path)
        self.document = None
        self.fields = []

    def load_document(self) -> bool:
        """
        Load the Word document.

        Returns:
            bool: True if document loaded successfully, False otherwise
        """
        try:
            self.document = Document(self.document_path)
            return True
        except Exception as e:
            print(f"Error loading document: {e}")
            return False

    def _collect_all_runs(self) -> List[Any]:
        """
        Collect all runs from the entire document, maintaining order.

        Returns:
            List[Any]: List of all XML run elements from the document
        """
        if not self.document:
            return []

        all_runs = []
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        # Collect runs from main document body
        for paragraph in self.document.paragraphs:
            paragraph_xml = paragraph._element
            runs = paragraph_xml.findall(".//w:r", ns)
            all_runs.extend(runs)

        # Collect runs from tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_xml = paragraph._element
                        runs = paragraph_xml.findall(".//w:r", ns)
                        all_runs.extend(runs)

        # Collect runs from headers and footers
        for section in self.document.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    paragraph_xml = paragraph._element
                    runs = paragraph_xml.findall(".//w:r", ns)
                    all_runs.extend(runs)

            if section.footer:
                for paragraph in section.footer.paragraphs:
                    paragraph_xml = paragraph._element
                    runs = paragraph_xml.findall(".//w:r", ns)
                    all_runs.extend(runs)

        return all_runs

    def _process_runs_for_fields(self, runs) -> List[Dict[str, Any]]:
        """
        Process a list of runs to extract fields that may span multiple paragraphs.

        Args:
            runs: List of XML run elements

        Returns:
            List[Dict[str, Any]]: List of extracted fields
        """
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        fields = []
        field_nesting_level = 0
        current_field_code = []
        field_result_parts = []
        in_field_code = False
        in_field_result = False

        for run in runs:
            # Check for field characters first
            fld_chars = run.findall(".//w:fldChar", ns)

            for fld_char in fld_chars:
                fld_char_type = fld_char.get(qn("w:fldCharType"))

                if fld_char_type == "begin":
                    field_nesting_level += 1
                    if field_nesting_level == 1:
                        # Start of top-level field
                        current_field_code = []
                        field_result_parts = []
                        in_field_code = True
                        in_field_result = False
                    else:
                        # Nested field - add opening brace
                        if in_field_code:
                            current_field_code.append("{")

                elif fld_char_type == "separate":
                    if field_nesting_level == 1:
                        # End of top-level field code, start of result
                        in_field_code = False
                        in_field_result = True

                elif fld_char_type == "end":
                    if field_nesting_level == 1:
                        # End of top-level field
                        field_code = "".join(current_field_code).strip()
                        field_result = "".join(field_result_parts).strip()

                        if field_code:
                            # Clean up field code - remove field results (text in « »)
                            cleaned_code = self._clean_field_code(field_code)

                            field_info = {
                                "type": (
                                    cleaned_code.split()[0].upper()
                                    if cleaned_code
                                    else "unknown"
                                ),
                                "code": cleaned_code,
                                "result": field_result,
                                "nested_fields": self._find_nested_fields_improved(
                                    cleaned_code
                                ),
                                "full_text": f"{{ {cleaned_code} }}",
                            }
                            fields.append(field_info)

                        in_field_code = False
                        in_field_result = False
                    else:
                        # End of nested field - add closing brace
                        if in_field_code:
                            current_field_code.append("}")

                    field_nesting_level -= 1

            # Process instruction text (field codes) - these are the actual field instructions
            if in_field_code:
                instr_texts = run.findall(".//w:instrText", ns)
                for instr_text in instr_texts:
                    if instr_text.text:
                        current_field_code.append(instr_text.text)

            # Process regular text - this includes both field code text and content text
            if in_field_code:
                text_elements = run.findall(".//w:t", ns)
                for text_elem in text_elements:
                    if text_elem.text:
                        current_field_code.append(text_elem.text)
            elif in_field_result:
                text_elements = run.findall(".//w:t", ns)
                for text_elem in text_elements:
                    if text_elem.text:
                        field_result_parts.append(text_elem.text)

        return fields

    def _clean_field_code(self, field_code: str) -> str:
        """
        Clean the field code by removing field results (text in « ») and fixing formatting.

        Args:
            field_code (str): The raw field code

        Returns:
            str: Cleaned field code
        """
        # Remove field results - text between « and »
        cleaned = re.sub(r"«[^»]*»", "", field_code)

        # Clean up extra spaces
        cleaned = re.sub(r"\s+", " ", cleaned)

        return cleaned.strip()

    def _find_nested_fields_improved(self, field_code: str) -> List[str]:
        """
        Improved method to find nested fields within a field code.

        Args:
            field_code (str): The field code to search for nested fields

        Returns:
            List[str]: List of nested field codes
        """
        nested_fields = []
        i = 0

        while i < len(field_code):
            if field_code[i] == "{":
                # Find matching closing brace
                brace_count = 1
                j = i + 1

                while j < len(field_code) and brace_count > 0:
                    if field_code[j] == "{":
                        brace_count += 1
                    elif field_code[j] == "}":
                        brace_count -= 1
                    j += 1

                if brace_count == 0:
                    # Found complete nested field
                    nested_field = field_code[i:j].strip()
                    if len(nested_field) > 2:  # More than just {}
                        # Clean up the nested field - remove extra spaces and field results
                        inner_content = nested_field[1:-1].strip()
                        if (
                            inner_content and not inner_content.isspace()
                        ):  # Only add if there's actual content
                            # Also clean the inner content to remove field results
                            cleaned_inner = self._clean_field_code(inner_content)
                            if (
                                cleaned_inner
                            ):  # Only add if there's still content after cleaning
                                nested_fields.append(f"{{ {cleaned_inner} }}")
                    i = j
                else:
                    i += 1
            else:
                i += 1

        return nested_fields

    def _find_nested_fields(self, field_code: str) -> List[str]:
        """
        Find nested fields within a field code.

        Args:
            field_code (str): The field code to search for nested fields

        Returns:
            List[str]: List of nested field codes
        """
        nested_fields = []

        # Look for patterns like { MERGEFIELD ... } within the field code
        i = 0
        while i < len(field_code):
            if field_code[i] == "{":
                # Find the matching closing brace
                brace_count = 1
                j = i + 1
                while j < len(field_code) and brace_count > 0:
                    if field_code[j] == "{":
                        brace_count += 1
                    elif field_code[j] == "}":
                        brace_count -= 1
                    j += 1

                if brace_count == 0:
                    nested_field = field_code[i:j].strip()
                    nested_fields.append(nested_field)
                    i = j
                else:
                    i += 1
            else:
                i += 1

        return nested_fields

    def extract_all_fields(self) -> List[Dict[str, Any]]:
        """
        Extract all fields from the document using document-level processing.
        This handles fields that span multiple paragraphs correctly.

        Returns:
            List[Dict[str, Any]]: List of all field dictionaries
        """
        if not self.document:
            if not self.load_document():
                return []

        # Use document-level processing to handle cross-paragraph fields
        all_runs = self._collect_all_runs()
        all_fields = self._process_runs_for_fields(all_runs)

        self.fields = all_fields
        return all_fields

    def print_fields_readable(self, fields: Optional[List[Dict[str, Any]]] = None):
        """
        Print all fields in a readable format.

        Args:
            fields (Optional[List[Dict[str, Any]]]): List of fields to print.
                                                   If None, uses self.fields
        """
        if fields is None:
            fields = self.fields

        if not fields:
            print("No fields found in the document.")
            return

        print(f"\n{'='*60}")
        print(f"WORD FIELDS FOUND IN: {self.document_path.name}")
        print(f"{'='*60}")
        print(f"Total fields found: {len(fields)}\n")

        for i, field in enumerate(fields, 1):
            print(f"Field #{i}")
            print(f"{'─'*40}")
            print(f"Type: {field['type']}")
            print(f"Code: {field['code']}")
            print(f"Full Text: {field['full_text']}")

            if field["result"]:
                print(f"Result: {field['result']}")

            if field["nested_fields"]:
                print(f"Nested Fields ({len(field['nested_fields'])}):")
                for j, nested in enumerate(field["nested_fields"], 1):
                    print(f"  {j}. {nested}")

            print()  # Empty line between fields


def extract_fields_from_document(document_path: str) -> List[Dict[str, Any]]:
    """
    Convenience function to extract fields from a document.

    Args:
        document_path (str): Path to the Word document

    Returns:
        List[Dict[str, Any]]: List of field dictionaries
    """
    extractor = WordFieldExtractor(document_path)
    return extractor.extract_all_fields()


def main():
    """
    Main function to run the field extractor.
    """
    # Set your document path here
    document_path = r"documents\Kodet version af brevkoder.docx"

    # Alternatively, you can use command line arguments
    if len(sys.argv) == 2:
        document_path = sys.argv[1]

    if not Path(document_path).exists():
        print(f"Error: Document '{document_path}' does not exist.")
        print("Available documents:")
        docs_folder = Path("documents")
        if docs_folder.exists():
            for doc in docs_folder.glob("*.docx"):
                if not doc.name.startswith("~$"):
                    print(f"  - {doc}")
        return

    print(f"Extracting fields from: {document_path}")

    extractor = WordFieldExtractor(document_path)
    fields = extractor.extract_all_fields()
    extractor.print_fields_readable()


if __name__ == "__main__":
    main()
